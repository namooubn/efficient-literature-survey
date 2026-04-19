#!/usr/bin/env python3
"""
Batch-extract metadata from a folder of literature files (PDF, DOCX, TXT, MD, EPUB).
Outputs a structured report with title, author, year, page/word count, text volume,
scanned-image detection (PDF only), and formatted citations.

Usage:
    python extract_literature_metadata.py /path/to/literature/folder

Dependencies:
    pip install PyPDF2 pdfplumber python-docx ebooklib beautifulsoup4
"""

import difflib
import hashlib
import json
import logging
import math
import os
import re
import sys
from concurrent.futures import ThreadPoolExecutor
from datetime import datetime
from pathlib import Path


# ---------------------------------------------------------------------------
# Logging
# ---------------------------------------------------------------------------

def _setup_logging(verbose: bool = False, quiet: bool = False) -> None:
    if quiet:
        level = logging.WARNING
    elif verbose:
        level = logging.DEBUG
    else:
        level = logging.INFO
    root = logging.getLogger()
    root.setLevel(level)
    if not root.handlers:
        logging.basicConfig(
            level=level,
            format="%(asctime)s [%(levelname)s] %(message)s",
            datefmt="%H:%M:%S",
        )


# ---------------------------------------------------------------------------
# Chinese compound surnames (复姓)
# ---------------------------------------------------------------------------

_COMPOUND_SURNAMES = {
    "欧阳", "司马", "诸葛", "上官", "皇甫", "令狐", "公孙", "慕容",
    "司徒", "司空", "长孙", "尉迟", "达奚", "赫连", "拓跋", "独孤",
    "完颜", "耶律", "宇文", "澹台", "贺兰", "纳兰", "呼延", "西门",
    "南宫", "东郭", "百里", "端木", "壤驷", "公良", "宰父", "谷梁",
}


def _split_chinese_name(name: str) -> tuple[str, str]:
    """Split a Chinese name into surname and given name, handling compound surnames."""
    name = name.strip().replace(" ", "")
    if not name:
        return ("", "")
    for cs in _COMPOUND_SURNAMES:
        if name.startswith(cs):
            return (cs, name[len(cs):])
    return (name[0], name[1:])


# ---------------------------------------------------------------------------
# Bibliographic metadata extraction from text
# ---------------------------------------------------------------------------

def _extract_bib_info_from_text(text: str) -> dict:
    """
    Attempt to extract journal/publisher/volume/issue/pages/year from the
    first page / front matter text using heuristic regexes.
    """
    info = {
        "journal": "",
        "publisher": "",
        "volume": "",
        "issue": "",
        "page_range": "",
        "doi": "",
    }
    if not text:
        return info

    # DOI
    doi_match = re.search(
        r'\b(10\.\d{4,}(?:\.\d+)*/[^\s"<>]+)', text, re.IGNORECASE
    )
    if doi_match:
        info["doi"] = doi_match.group(1)

    # Volume / Issue – English patterns
    vol_match = re.search(
        r'(?:Vol\.?|Volume)\s*(\d+)(?:\s*[,:]?\s*(?:No\.?|Issue|#)\s*(\d+))?', text, re.IGNORECASE
    )
    if vol_match:
        info["volume"] = vol_match.group(1)
        if vol_match.group(2):
            info["issue"] = vol_match.group(2)

    # Pages – English patterns: pp. 123-145, p. 123, 123–145
    page_match = re.search(r'(?:pp?\.?\s*)?(\d+)[\s–—-]+(\d+)', text)
    if page_match:
        info["page_range"] = f"{page_match.group(1)}-{page_match.group(2)}"

    # Volume / Issue / Pages – Chinese patterns
    cn_vol = re.search(r'第\s*(\d+)\s*卷', text)
    cn_issue = re.search(r'第\s*(\d+)\s*期', text)
    cn_page = re.search(r'第\s*(\d+)\s*[—–~\-]\s*(\d+)\s*页', text)
    if cn_vol:
        info["volume"] = cn_vol.group(1)
    if cn_issue:
        info["issue"] = cn_issue.group(1)
    if cn_page:
        info["page_range"] = f"{cn_page.group(1)}-{cn_page.group(2)}"

    # Journal / Publisher heuristics
    # Common Chinese journal markers
    cn_journal = re.search(r'《([^》]{3,50})》', text)
    if cn_journal:
        info["journal"] = cn_journal.group(1).strip()

    # Try to grab a line that looks like a journal header (all caps, short)
    for line in text.splitlines()[:15]:
        line = line.strip()
        if 10 < len(line) < 80 and line.isupper() and line.replace(" ", "").isalpha():
            if not info["journal"]:
                info["journal"] = line.title()
            break

    return info


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _estimate_pages_from_chars(char_count: int, chars_per_page: int = 500) -> int:
    """Estimate page count for text-based files where true pagination is unknown."""
    return max(1, math.ceil(char_count / chars_per_page))


def _safe_truncate(text: str, limit: int = 2000) -> str:
    return text[:limit] if text else ""


def _smart_word_count(text: str) -> int:
    """
    Estimate word count for mixed Chinese/English text.
    Chinese: count characters (excluding whitespace and punctuation).
    English: count space-separated tokens.
    """
    if not text:
        return 0
    cn_chars = re.findall(r'[\u4e00-\u9fff\u3400-\u4dbf\U00020000-\U0002a6df\U0002a700-\U0002b73f]', text)
    en_tokens = re.findall(r'\b[a-zA-Z]+\b', text)
    if len(cn_chars) == 0:
        return len(text.split())
    return len(cn_chars) + len(en_tokens)


def _detect_scanned_pdf(read_page_texts: list[str], total_pages: int) -> bool:
    """
    Multi-page sampling heuristic for scanned/image-based PDF detection.
    Only considers pages that were actually read (read_page_texts), avoiding
    false positives from unread pages that default to empty strings.
    """
    if total_pages <= 1 or not read_page_texts:
        return False

    total_chars = sum(len(t) for t in read_page_texts)

    if total_chars > 8000:
        return False

    read_count = len(read_page_texts)

    indices = [0]
    if read_count > 3:
        mid = read_count // 2
        if mid > 0:
            indices.append(mid)
    if read_count > 5:
        end = read_count - 1
        if end not in indices:
            indices.append(end)

    sampled = [read_page_texts[i] for i in indices if i < read_count]
    empty_samples = sum(1 for t in sampled if len(t.strip()) < 25)

    if len(sampled) >= 2 and empty_samples / len(sampled) >= 0.66:
        return True

    density = total_chars / read_count if read_count > 0 else 0
    if density < 12 and read_count > 3:
        return True

    if total_chars < 50:
        return True

    return False


def _get_pdf_read_pages(total_pages: int, max_pages: int = 0) -> list[int]:
    """
    Decide which pages to read based on document length.

    - Short docs (<= 50 pages): read all pages.
    - Long docs (> 50 pages): read first 15 + last 5 pages to avoid
      getting stuck in front-matter (cover, TOC, preface) while still
      sampling the body and conclusion.
    - max_pages > 0 overrides the heuristic and caps total pages read.
    """
    if max_pages > 0:
        return list(range(min(total_pages, max_pages)))
    if total_pages <= 50:
        return list(range(total_pages))
    pages = list(range(min(15, total_pages)))
    pages.extend(range(max(0, total_pages - 5), total_pages))
    return sorted(set(pages))


# ---------------------------------------------------------------------------
# Duplicate detection
# ---------------------------------------------------------------------------


def _find_duplicates(results: list[dict], threshold: float = 0.80) -> dict[str, list[str]]:
    """
    Detect duplicate references by title similarity (SequenceMatcher).
    Returns {filename: [list of duplicate filenames]}.
    """
    dups: dict[str, list[str]] = {}
    n = len(results)
    for i in range(n):
        title_i = (results[i].get("title_from_meta") or results[i]["filename"]).lower().strip()
        if not title_i:
            continue
        for j in range(i + 1, n):
            title_j = (results[j].get("title_from_meta") or results[j]["filename"]).lower().strip()
            if not title_j:
                continue
            ratio = difflib.SequenceMatcher(None, title_i, title_j).ratio()
            if ratio >= threshold:
                fi = results[i]["filename"]
                fj = results[j]["filename"]
                dups.setdefault(fi, []).append(f"{fj} (相似度 {ratio:.0%})")
                dups.setdefault(fj, []).append(f"{fi} (相似度 {ratio:.0%})")
    for k in dups:
        seen = set()
        uniq = []
        for v in dups[k]:
            if v not in seen:
                seen.add(v)
                uniq.append(v)
        dups[k] = uniq
    return dups


# ---------------------------------------------------------------------------
# Citation generation
# ---------------------------------------------------------------------------

def _normalize_author(author_raw: str) -> list[str]:
    """Split author string into individual names."""
    if not author_raw:
        return []
    parts = re.split(r"[,;，；/&]", author_raw)
    return [p.strip() for p in parts if p.strip()]


def _guess_doc_type(metadata: dict) -> str:
    """Guess document type code for GB/T 7714."""
    fmt = metadata.get("format", "").lower()
    pages = metadata.get("pages", 0)
    # Heuristic: >200 pages likely monograph/book; otherwise article
    if pages > 200:
        return "M"  # Monograph
    if fmt == "epub":
        return "M"
    return "J"  # Default to journal article


def generate_citation(metadata: dict, style: str = "numbered") -> str:
    """
    Generate a formatted citation string from metadata.

    Supported styles:
      - "gb7714"   : GB/T 7714 Chinese standard
      - "apa"      : APA 7th edition (simplified)
      - "mla"      : MLA 9th edition (simplified)
      - "numbered" : [N] numeric citation

    Missing fields are omitted (not fabricated).  When critical fields
    are missing, a trailing notice is appended.
    """
    title = (metadata.get("title_from_meta") or metadata.get("filename", "")).strip()
    author_raw = metadata.get("author_from_meta", "").strip()
    authors = _normalize_author(author_raw)
    year = metadata.get("year", "")
    journal = metadata.get("journal", "")
    publisher = metadata.get("publisher", "")
    volume = metadata.get("volume", "")
    issue = metadata.get("issue", "")
    page_range = metadata.get("page_range", "")
    doi = metadata.get("doi", "")
    doc_type = _guess_doc_type(metadata)

    incomplete = False
    # Critical missing fields check
    if not year or (not journal and not publisher and doc_type == "J"):
        incomplete = True

    def _author_str_gb(authors_list: list[str]) -> str:
        if not authors_list:
            return "佚名"
        if len(authors_list) == 1:
            return authors_list[0]
        if len(authors_list) == 2:
            return f"{authors_list[0]}, {authors_list[1]}"
        if len(authors_list) == 3:
            return f"{authors_list[0]}, {authors_list[1]}, {authors_list[2]}"
        return f"{authors_list[0]} 等"

    def _author_str_apa(authors_list: list[str]) -> str:
        if not authors_list:
            return "Anonymous"
        if len(authors_list) == 1:
            return _apa_name(authors_list[0])
        if len(authors_list) == 2:
            return f"{_apa_name(authors_list[0])} & {_apa_name(authors_list[1])}"
        return f"{_apa_name(authors_list[0])} et al."

    def _author_str_mla(authors_list: list[str]) -> str:
        if not authors_list:
            return "Anonymous"
        first = _mla_name(authors_list[0])
        if len(authors_list) > 1:
            return f"{first}, et al."
        return first

    parts: list[str] = []

    if style == "gb7714":
        author_str = _author_str_gb(authors)
        parts.append(f"[{author_str}]. {title}[{doc_type}]")
        if publisher:
            parts.append(publisher)
        elif journal:
            parts.append(journal)
        if year:
            parts.append(year)
        if volume:
            vi = f"{volume}"
            if issue:
                vi += f"({issue})"
            parts.append(vi)
        if page_range:
            parts.append(f": {page_range}")
        if doi:
            parts.append(f". DOI:{doi}")
        citation = ". ".join(parts)
        if citation.endswith("."):
            citation = citation[:-1]
        citation += "."

    elif style == "apa":
        author_str = _author_str_apa(authors)
        year_part = f" ({year})" if year else ""
        citation = f"{author_str}{year_part}. {title}."
        if journal:
            citation += f" *{journal}*"
            if volume:
                citation += f", *{volume}*"
                if issue:
                    citation += f"({issue})"
            if page_range:
                citation += f", {page_range}"
            citation += "."
        elif publisher:
            citation += f" {publisher}."
        if doi:
            citation += f" https://doi.org/{doi}"

    elif style == "mla":
        author_str = _author_str_mla(authors)
        citation = f'"{title}." '
        if journal:
            citation += f"*{journal}*, "
            if volume:
                citation += f"vol. {volume}, "
            if issue:
                citation += f"no. {issue}, "
        if year:
            citation += f"{year}, "
        if page_range:
            citation += f"pp. {page_range}."
        else:
            citation = citation.rstrip(", ") + "."
        citation = f"{author_str}. {citation}"

    else:  # numbered
        num = metadata.get("_citation_number", "")
        prefix = f"[{num}] " if num else ""
        parts: list[str] = []
        if authors:
            parts.append(authors[0] + (" 等" if len(authors) > 1 else ""))
        parts.append(title)
        if journal:
            parts.append(f"《{journal}》")
        elif publisher:
            parts.append(publisher)
        if year:
            parts.append(year)
        citation = prefix + ", ".join([p for p in parts if p])
        citation += "."

    if incomplete:
        citation += " 〔文献信息不完整，请手动补全〕"
    return citation


def _apa_name(name: str) -> str:
    """Convert 'Zhang San' or 'San Zhang' to APA-style 'Zhang, S.' heuristic."""
    name = name.strip()
    if any('\u4e00' <= c <= '\u9fff' for c in name):
        surname, given = _split_chinese_name(name)
        if surname:
            given_initials = "".join(f"{g[0]}." for g in given if g)
            return f"{surname}, {given_initials}"
    parts = name.split()
    if len(parts) >= 2 and len(parts[-1]) == 1:
        return f"{parts[-1]}, {''.join(p[0] + '.' for p in parts[:-1])}"
    if len(parts) >= 2:
        return f"{parts[-1]}, {parts[0][0]}."
    return name


def _mla_name(name: str) -> str:
    """Convert name to MLA-style 'Last, First' heuristic."""
    name = name.strip()
    if any('\u4e00' <= c <= '\u9fff' for c in name):
        surname, given = _split_chinese_name(name)
        if surname:
            return f"{surname}, {given}"
    parts = name.split()
    if len(parts) >= 2:
        return f"{parts[-1]}, {parts[0]}"
    return name


# ---------------------------------------------------------------------------
# BibTeX generation
# ---------------------------------------------------------------------------

def _guess_bibtex_type(metadata: dict) -> str:
    """Map metadata to a BibTeX entry type."""
    fmt = metadata.get("format", "").lower()
    pages = metadata.get("pages", 0)
    journal = metadata.get("journal", "")
    if journal:
        return "article"
    if pages > 200 or fmt == "epub":
        return "book"
    return "misc"


def _bibtex_escape(s: str) -> str:
    return s.replace("{", "\\{").replace("}", "\\}").replace("&", "\\&")


def generate_bibtex(results: list[dict]) -> str:
    """Generate a .bib file from extracted metadata."""
    lines: list[str] = []
    for idx, r in enumerate(results, start=1):
        entry_type = _guess_bibtex_type(r)
        citekey = re.sub(r"[^a-zA-Z0-9]", "", (r.get("author_from_meta") or "Unknown").split(",")[0].split(" ")[0])
        citekey = citekey or "Unknown"
        year = r.get("year", "")
        citekey = f"{citekey}{year}{idx}"
        lines.append(f"@{entry_type}{{{citekey},")
        lines.append(f"  title = {{{_bibtex_escape(r.get('title_from_meta', r['filename']))}}},")
        authors = _normalize_author(r.get("author_from_meta", ""))
        if authors:
            lines.append(f"  author = {{{_bibtex_escape(' and '.join(authors))}}},")
        if year:
            lines.append(f"  year = {{{year}}},")
        if r.get("journal"):
            lines.append(f"  journal = {{{_bibtex_escape(r['journal'])}}},")
        if r.get("publisher"):
            lines.append(f"  publisher = {{{_bibtex_escape(r['publisher'])}}},")
        if r.get("volume"):
            lines.append(f"  volume = {{{r['volume']}}},")
        if r.get("issue"):
            lines.append(f"  number = {{{r['issue']}}},")
        if r.get("page_range"):
            lines.append(f"  pages = {{{r['page_range']}}},")
        if r.get("doi"):
            lines.append(f"  doi = {{{r['doi']}}},")
        if not r.get("journal") and not r.get("publisher"):
            lines.append("  note = {文献信息不完整，请手动补全},")
        lines.append("}")
        lines.append("")
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Incremental / caching helpers
# ---------------------------------------------------------------------------

def _file_sha256(file_path: str) -> str:
    """Compute SHA-256 hash of a file for cache invalidation."""
    h = hashlib.sha256()
    try:
        with open(file_path, "rb") as f:
            while chunk := f.read(8192):
                h.update(chunk)
    except Exception:
        return ""
    return h.hexdigest()


def _load_cache(cache_path: Path) -> dict:
    """Load cached extraction results keyed by filename -> {sha256, result}."""
    if not cache_path.exists():
        return {}
    try:
        with open(cache_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        if isinstance(data, dict) and "files" in data:
            return data["files"]
    except Exception:
        pass
    return {}


def _save_cache(cache_path: Path, cache_data: dict) -> None:
    """Save cache with version and timestamp metadata."""
    payload = {
        "version": 1,
        "generated_at": datetime.now().isoformat(),
        "files": cache_data,
    }
    try:
        with open(cache_path, "w", encoding="utf-8") as f:
            json.dump(payload, f, ensure_ascii=False, indent=2)
    except Exception as e:
        logging.warning("缓存保存失败：%s", e)


# ---------------------------------------------------------------------------
# Interactive CLI helpers
# ---------------------------------------------------------------------------

def _prompt_path() -> Path:
    """Prompt user for a folder path with validation."""
    while True:
        raw = input("\n请输入文献文件夹路径（例如 /home/alice/references）：\n> ").strip()
        p = Path(raw).expanduser()
        if not p.exists():
            print(f"  路径不存在：{p}")
            continue
        if not p.is_dir():
            print(f"  这不是一个文件夹：{p}")
            continue
        return p


def _interactive_prompt_unsupported(skipped: list, supported_exts: set) -> None:
    """Explain why files were skipped in a user-friendly way."""
    if not skipped:
        return
    print(f"\n  检测到 {len(skipped)} 个不支持的文件：")
    for name in skipped:
        ext = Path(name).suffix.lower()
        print(f"    - {name} ({ext or '无后缀'})")
    print(f"  当前支持的格式：{', '.join(sorted(supported_exts))}")


# ---------------------------------------------------------------------------
# Extractors by format
# ---------------------------------------------------------------------------

def extract_pdf(pdf_path: str, max_pages: int = 0) -> dict:
    """Extract metadata and text sample from a single PDF."""
    result = {
        "filename": os.path.basename(pdf_path),
        "relative_path": "",
        "format": "pdf",
        "pages": 0,
        "word_count": 0,
        "text_chars": 0,
        "first_page_text": "",
        "is_scanned": False,
        "author_from_meta": "",
        "title_from_meta": "",
        "year": "",
        "journal": "",
        "publisher": "",
        "volume": "",
        "issue": "",
        "page_range": "",
        "doi": "",
        "note": "",
    }

    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(pdf_path)
        result["pages"] = len(reader.pages)
        meta = reader.metadata
        if meta:
            result["author_from_meta"] = str(meta.get("/Author", ""))
            result["title_from_meta"] = str(meta.get("/Title", ""))
            creation_date = str(meta.get("/CreationDate", ""))
            m = re.search(r'D:(\d{4})', creation_date)
            if m:
                result["year"] = m.group(1)

        read_pages = _get_pdf_read_pages(result["pages"], max_pages)
        read_page_texts: list[str] = []
        total_text = ""
        for i in read_pages:
            try:
                page = reader.pages[i]
                txt = page.extract_text() or ""
                read_page_texts.append(txt)
                total_text += txt
                if i == 0:
                    result["first_page_text"] = _safe_truncate(txt)
            except Exception as e:
                if not result["note"]:
                    result["note"] = ""
                result["note"] += f"Page {i} extract error: {type(e).__name__}; "
        result["text_chars"] = len(total_text)
    except Exception as e:
        read_page_texts = []
        if not result["note"]:
            result["note"] = ""
        result["note"] += f"PyPDF2 error: {type(e).__name__}: {e}; "

    # Fallback to pdfplumber if PyPDF2 extracted very little text
    if result["text_chars"] < 500 and result["pages"] > 0:
        try:
            import pdfplumber
            with pdfplumber.open(pdf_path) as pdf:
                result["pages"] = len(pdf.pages)
                read_pages = _get_pdf_read_pages(result["pages"], max_pages)
                read_page_texts = []
                total_text = ""
                for i in read_pages:
                    txt = pdf.pages[i].extract_text() or ""
                    read_page_texts.append(txt)
                    total_text += txt
                    if i == 0:
                        result["first_page_text"] = _safe_truncate(txt)
                result["text_chars"] = len(total_text)
        except Exception as e:
            if not result["note"]:
                result["note"] = ""
            result["note"] += f"pdfplumber fallback error: {type(e).__name__}: {e}; "

    result["word_count"] = _smart_word_count(total_text) if result["text_chars"] > 0 else 0

    # Extract bibliographic info from first-page text
    bib_info = _extract_bib_info_from_text(result["first_page_text"])
    result.update(bib_info)

    # Fallback: extract year from filename if not found in metadata
    if not result.get("year"):
        m = re.search(r'\b(19|20)\d{2}\b', result["filename"])
        if m:
            result["year"] = m.group(0)

    result["is_scanned"] = _detect_scanned_pdf(read_page_texts, result["pages"])

    return result


def extract_docx(docx_path: str) -> dict:
    """Extract metadata and text sample from a DOCX file."""
    result = {
        "filename": os.path.basename(docx_path),
        "relative_path": "",
        "format": "docx",
        "pages": 0,
        "word_count": 0,
        "text_chars": 0,
        "first_page_text": "",
        "is_scanned": False,
        "author_from_meta": "",
        "title_from_meta": "",
        "year": "",
        "journal": "",
        "publisher": "",
        "volume": "",
        "issue": "",
        "page_range": "",
        "doi": "",
        "note": "",
    }

    try:
        from docx import Document
        doc = Document(docx_path)

        core_props = doc.core_properties
        result["author_from_meta"] = core_props.author or ""
        result["title_from_meta"] = core_props.title or ""

        paragraphs = []
        char_count = 0
        for p in doc.paragraphs:
            paragraphs.append(p)
            char_count += len(p.text)
            if char_count >= 5000:
                break
        full_text = "\n".join(p.text for p in paragraphs)
        result["text_chars"] = len(full_text)
        result["word_count"] = _smart_word_count(full_text)
        result["first_page_text"] = _safe_truncate(full_text)
        result["pages"] = _estimate_pages_from_chars(result["text_chars"])

        table_texts = []
        for table in doc.tables[:5]:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        table_texts.append(cell.text.strip())
        if table_texts:
            table_full = "\n".join(table_texts)
            result["word_count"] += _smart_word_count(table_full)

        # Try to extract bibliographic info from front matter
        bib_info = _extract_bib_info_from_text(full_text)
        result.update(bib_info)

    except Exception as e:
        result["note"] = f"DOCX extraction error: {e}"

    return result


def extract_txt(txt_path: str) -> dict:
    """Extract text sample from a plain-text file (TXT or MD)."""
    result = {
        "filename": os.path.basename(txt_path),
        "relative_path": "",
        "format": Path(txt_path).suffix.lower().lstrip("."),
        "pages": 0,
        "word_count": 0,
        "text_chars": 0,
        "first_page_text": "",
        "is_scanned": False,
        "author_from_meta": "",
        "title_from_meta": "",
        "year": "",
        "journal": "",
        "publisher": "",
        "volume": "",
        "issue": "",
        "page_range": "",
        "doi": "",
        "note": "",
    }

    try:
        with open(txt_path, "r", encoding="utf-8") as f:
            content = f.read()
    except UnicodeDecodeError:
        try:
            with open(txt_path, "r", encoding="gbk") as f:
                content = f.read()
        except Exception as e:
            result["note"] = f"Encoding error: {e}"
            return result
    except Exception as e:
        result["note"] = f"Read error: {e}"
        return result

    result["text_chars"] = len(content)
    result["word_count"] = _smart_word_count(content)
    result["first_page_text"] = _safe_truncate(content)
    result["pages"] = _estimate_pages_from_chars(result["text_chars"])

    # Try to extract bibliographic info
    bib_info = _extract_bib_info_from_text(content[:3000])
    result.update(bib_info)

    return result


def extract_epub(epub_path: str) -> dict:
    """Extract metadata and text sample from an EPUB file."""
    result = {
        "filename": os.path.basename(epub_path),
        "relative_path": "",
        "format": "epub",
        "pages": 0,
        "word_count": 0,
        "text_chars": 0,
        "first_page_text": "",
        "is_scanned": False,
        "author_from_meta": "",
        "title_from_meta": "",
        "year": "",
        "journal": "",
        "publisher": "",
        "volume": "",
        "issue": "",
        "page_range": "",
        "doi": "",
        "note": "",
    }

    try:
        import ebooklib
        from ebooklib import epub
        from bs4 import BeautifulSoup

        book = epub.read_epub(epub_path)

        title_list = book.get_metadata("DC", "title")
        if title_list:
            result["title_from_meta"] = str(title_list[0][0])
        creator_list = book.get_metadata("DC", "creator")
        if creator_list:
            result["author_from_meta"] = str(creator_list[0][0])

        all_texts = []
        count = 0
        for item in book.get_items():
            if item.get_type() == ebooklib.ITEM_DOCUMENT:
                soup = BeautifulSoup(item.get_content(), "html.parser")
                text = soup.get_text(separator="\n", strip=True)
                if text:
                    all_texts.append(text)
                    count += 1
                if count >= 10:
                    break

        full_text = "\n".join(all_texts)
        result["text_chars"] = len(full_text)
        result["word_count"] = _smart_word_count(full_text)
        result["first_page_text"] = _safe_truncate(full_text)
        result["pages"] = _estimate_pages_from_chars(result["text_chars"])

        bib_info = _extract_bib_info_from_text(full_text[:3000])
        result.update(bib_info)

    except ImportError:
        result["note"] = "Missing dependency: pip install ebooklib beautifulsoup4"
    except Exception as e:
        result["note"] = f"EPUB extraction error: {e}"

    return result


# ---------------------------------------------------------------------------
# Dispatcher
# ---------------------------------------------------------------------------

SUPPORTED_EXTS = {".pdf", ".docx", ".txt", ".md", ".epub"}
CAJ_EXT = {".caj"}


def extract_info(file_path: str, max_pages: int = 0) -> dict:
    """Dispatch to the correct extractor based on file extension."""
    ext = Path(file_path).suffix.lower()

    if ext == ".pdf":
        return extract_pdf(file_path, max_pages=max_pages)
    if ext == ".docx":
        return extract_docx(file_path)
    if ext in (".txt", ".md"):
        return extract_txt(file_path)
    if ext == ".epub":
        return extract_epub(file_path)

    if ext in CAJ_EXT:
        return {
            "filename": os.path.basename(file_path),
            "relative_path": "",
            "format": "caj",
            "pages": 0,
            "word_count": 0,
            "text_chars": 0,
            "first_page_text": "",
            "is_scanned": False,
            "author_from_meta": "",
            "title_from_meta": "",
            "year": "",
            "journal": "",
            "publisher": "",
            "volume": "",
            "issue": "",
            "page_range": "",
            "doi": "",
            "note": "CAJ format is not directly supported. Please convert to PDF using CAJViewer or caj2pdf first.",
        }

    return {
        "filename": os.path.basename(file_path),
        "relative_path": "",
        "format": ext.lstrip("."),
        "pages": 0,
        "word_count": 0,
        "text_chars": 0,
        "first_page_text": "",
        "is_scanned": False,
        "author_from_meta": "",
        "title_from_meta": "",
        "year": "",
        "journal": "",
        "publisher": "",
        "volume": "",
        "issue": "",
        "page_range": "",
        "doi": "",
        "note": f"Unsupported file format: {ext}",
    }


# ---------------------------------------------------------------------------
# Markdown report generator
# ---------------------------------------------------------------------------

_OCR_GUIDANCE = """
> **处理建议**：以下工具可将其转换为可提取文本的 PDF
> - **marker**（推荐，GPU 快、排版保留好）：`pip install marker-pdf && marker_single <文件>`
> - **nougat**（学术 PDF 专用）：`pip install nougat-ocr && nougat <文件>`
> - **pdf2image + pytesseract**（CPU 可用）：`pip install pdf2image pytesseract`，然后逐页 OCR
> - 转换完成后，重新运行本脚本以获取完整元数据。
"""


def generate_markdown_report(
    results: list,
    lit_dir: Path,
    skipped: list,
    duplicates: dict | None = None,
    citation_style: str = "gb7714",
) -> str:
    """Generate a human-readable Markdown report from extraction results."""
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    scanned = [r for r in results if r["is_scanned"]]
    by_format = {}
    by_subdir: dict[str, list] = {}
    for r in results:
        fmt = r["format"]
        by_format.setdefault(fmt, []).append(r)
        subdir = str(Path(r.get("relative_path", "")).parent)
        if subdir and subdir != ".":
            by_subdir.setdefault(subdir, []).append(r)

    total_pages = sum(r["pages"] for r in results)
    total_words = sum(r["word_count"] for r in results)
    dup_count = len(duplicates) if duplicates else 0

    lines = [
        "# 文献提取报告\n",
        f"**生成时间**：{now}",
        f"**文献目录**：`{lit_dir}`\n",
        "## 汇总统计\n",
        f"| 指标 | 数值 |",
        f"|------|------|",
        f"| 总文献数 | {len(results)} 篇 |",
        f"| 总页数（含估算） | {total_pages} 页 |",
        f"| 总字数（含估算） | {total_words} 字 |",
        f"| 扫描件（需 OCR） | {len(scanned)} 篇 |",
        f"| 疑似重复 | {dup_count} 篇 |",
        f"| 格式分布 | {', '.join(f'{k}: {len(v)} 篇' for k, v in by_format.items())} |",
        "",
    ]

    if scanned:
        lines.extend([
            "## ⚠️ 扫描件清单（需 OCR 处理）\n",
            "| 文件名 | 页数 | 状态 |",
            "|--------|------|------|",
        ])
        for r in scanned:
            lines.append(f"| {r['filename']} | {r['pages']} | 扫描 PDF，无法直接提取全文 |")
        lines.append("")
        lines.append(_OCR_GUIDANCE)
        lines.append("")

    if duplicates:
        lines.extend([
            "## 🔁 疑似重复文献\n",
            "| 文件名 | 疑似重复对象 |",
            "|--------|-------------|",
        ])
        for fname, dlist in sorted(duplicates.items()):
            lines.append(f"| {fname} | {'; '.join(dlist)} |")
        lines.append("")

    # Group by subfolder if any
    if by_subdir:
        lines.extend(["## 按子文件夹分组\n"])
        for subdir, items in sorted(by_subdir.items()):
            lines.append(f"### {subdir} ({len(items)} 篇)\n")
            lines.append("| 文件名 | 格式 | 页数 | 字数 | 作者 | 标题 | 备注 |")
            lines.append("|--------|------|------|------|------|------|------|")
            for r in sorted(items, key=lambda x: x["filename"].lower()):
                author = r["author_from_meta"] or "—"
                title = r["title_from_meta"] or "—"
                note = r["note"] or "—"
                if len(title) > 40:
                    title = title[:37] + "..."
                if len(author) > 30:
                    author = author[:27] + "..."
                lines.append(
                    f"| {r['filename']} | {r['format']} | {r['pages']} | {r['word_count']} | "
                    f"{author} | {title} | {note} |"
                )
            lines.append("")

    lines.extend([
        "## 文献详情（按文件名排序）\n",
        "| 序号 | 文件名 | 格式 | 页数 | 字数 | 作者 | 标题 | 备注 |",
        "|------|--------|------|------|------|------|------|------|",
    ])
    for r in sorted(results, key=lambda x: x["filename"].lower()):
        author = r["author_from_meta"] or "—"
        title = r["title_from_meta"] or "—"
        note = r["note"] or "—"
        if len(title) > 40:
            title = title[:37] + "..."
        if len(author) > 30:
            author = author[:27] + "..."
        num = r.get("_citation_number", "")
        lines.append(
            f"| {num} | {r['filename']} | {r['format']} | {r['pages']} | {r['word_count']} | "
            f"{author} | {title} | {note} |"
        )
    lines.append("")

    lines.extend([
        f"## 参考文献列表（{citation_style.upper()} 格式）\n",
    ])
    for r in sorted(results, key=lambda x: x.get("_citation_number", 0)):
        citation = r.get("citation", "")
        if citation:
            lines.append(f"{r.get('_citation_number', '')}. {citation}")
    lines.append("")

    if skipped:
        lines.extend([
            "## 跳过的文件\n",
            "| 文件名 | 原因 |",
            "|--------|------|",
        ])
        for name in skipped:
            ext = Path(name).suffix.lower()
            reason = f"不支持的格式 {ext}" if ext else "无文件后缀"
            lines.append(f"| {name} | {reason} |")
        lines.append("")

    lines.append("---\n*由 efficient-literature-survey 自动生成*\n")
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    args = sys.argv[1:]

    # Help
    if any(a in args for a in ("-h", "--help")):
        print(__doc__)
        sys.exit(0)

    # Parse optional flags
    max_pages = 0
    citation_style = "gb7714"
    output_dir: Path | None = None
    bibtex = False
    verbose = False
    quiet = False
    recursive = True
    i = 0
    while i < len(args):
        if args[i] in ("--max-pages", "-m") and i + 1 < len(args):
            try:
                max_pages = int(args[i + 1])
                if max_pages < 0:
                    max_pages = 0
            except ValueError:
                pass
            i += 2
        elif args[i] in ("--citation-style", "-c") and i + 1 < len(args):
            citation_style = args[i + 1].lower()
            i += 2
        elif args[i] in ("--output-dir", "-o") and i + 1 < len(args):
            output_dir = Path(args[i + 1]).expanduser()
            i += 2
        elif args[i] in ("--bibtex", "-b"):
            bibtex = True
            i += 1
        elif args[i] in ("--verbose", "-v"):
            verbose = True
            i += 1
        elif args[i] in ("--quiet", "-q"):
            quiet = True
            i += 1
        elif args[i] in ("--no-recursive",):
            recursive = False
            i += 1
        else:
            i += 1

    _setup_logging(verbose=verbose, quiet=quiet)

    # Determine folder path (first non-flag positional arg)
    pos_args = [a for a in args if not a.startswith("-")]
    if not pos_args:
        print("用法：python extract_literature_metadata.py <文献文件夹路径>")
        print("  [--max-pages N] [--citation-style gb7714|apa|mla|numbered]")
        print("  [--output-dir PATH] [--bibtex] [--verbose] [--quiet] [--no-recursive]")
        print("未提供路径，进入交互模式...")
        lit_dir = _prompt_path()
    else:
        lit_dir = Path(pos_args[0]).expanduser()
        if not lit_dir.exists():
            print(f"路径不存在：{lit_dir}")
            print("进入交互模式...")
            lit_dir = _prompt_path()
        elif not lit_dir.is_dir():
            print(f"这不是一个文件夹：{lit_dir}")
            print("进入交互模式...")
            lit_dir = _prompt_path()

    # Collect files (recursive by default)
    if recursive:
        all_files = sorted([f for f in lit_dir.rglob("*") if f.is_file()])
    else:
        all_files = sorted([f for f in lit_dir.iterdir() if f.is_file()])

    lit_files = [f for f in all_files if f.suffix.lower() in SUPPORTED_EXTS | CAJ_EXT]
    skipped = [f.name for f in all_files if f.suffix.lower() not in SUPPORTED_EXTS | CAJ_EXT]

    if not lit_files:
        logging.error("在 %s 中未找到支持的文献文件。", lit_dir)
        _interactive_prompt_unsupported(skipped, SUPPORTED_EXTS | CAJ_EXT)
        sys.exit(1)

    logging.info("找到 %d 篇文献，开始处理...", len(lit_files))
    if skipped:
        logging.info("跳过 %d 个不支持的文件", len(skipped))

    # Determine output directory
    if output_dir is None:
        output_dir = lit_dir / ".els_output"
    output_dir.mkdir(parents=True, exist_ok=True)

    # Load cache for incremental extraction
    cache_path = output_dir / "_literature_cache.json"
    cache = _load_cache(cache_path)
    results: list[dict] = []
    to_extract: list[Path] = []
    cache_hits = 0

    for lf in lit_files:
        file_hash = _file_sha256(str(lf))
        cached = cache.get(lf.name)
        if cached and cached.get("sha256") == file_hash and cached.get("result"):
            results.append(cached["result"])
            cache_hits += 1
        else:
            to_extract.append(lf)

    if cache_hits:
        logging.info("缓存命中 %d 篇（文件未变更），跳过重新提取", cache_hits)
    if to_extract:
        logging.info("需要重新提取 %d 篇", len(to_extract))

    # Concurrent extraction with ThreadPoolExecutor (I/O-bound: reading files)
    if to_extract:
        max_workers = min(4, len(to_extract)) if to_extract else 1
        logging.info("使用 %d 线程并发处理...", max_workers)

        def _extract_with_log(path: str) -> dict:
            info = extract_info(path, max_pages=max_pages)
            rel = Path(path).relative_to(lit_dir)
            info["relative_path"] = str(rel)
            info["filename"] = rel.name
            logging.info(
                "  [DONE] %s P%d WC%d | %s",
                info["format"], info["pages"], info["word_count"], str(rel)[:50],
            )
            return info

        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            extracted = list(executor.map(_extract_with_log, (str(lf) for lf in to_extract)))
        results.extend(extracted)

    # Also set relative_path for cached results if missing
    for r in results:
        if not r.get("relative_path"):
            for lf in lit_files:
                if lf.name == r.get("filename"):
                    r["relative_path"] = str(lf.relative_to(lit_dir))
                    break

    # Update cache with fresh results
    new_cache: dict[str, dict] = {}
    for r in results:
        fname = r.get("filename", "")
        for lf in lit_files:
            if lf.name == fname:
                new_cache[fname] = {
                    "sha256": _file_sha256(str(lf)),
                    "result": r,
                }
                break
    _save_cache(cache_path, new_cache)

    # Duplicate detection
    duplicates = _find_duplicates(results)
    if duplicates:
        logging.info("检测到 %d 篇疑似重复文献", len(duplicates))
        for fname, dlist in duplicates.items():
            logging.info("  - %s ↔ %s", fname, ", ".join(dlist))

    # Generate citations for each result
    for idx, r in enumerate(results, start=1):
        r["_citation_number"] = idx
        r["citation"] = generate_citation(r, citation_style)

    # Console summary
    print("\n" + "=" * 80)
    print("SUMMARY")
    print("=" * 80)
    for r in results:
        status = "⚠️ SCANNED" if r["is_scanned"] else "  OK     "
        dup_mark = " [DUP]" if r["filename"] in duplicates else ""
        note = f" | NOTE: {r['note']}" if r["note"] else ""
        rel = r.get("relative_path", r["filename"])
        print(
            f"[{status}] {r['format']:5} P{r['pages']:3} WC{r['word_count']:6} | "
            f"{rel[:45]}{dup_mark}{note}"
        )

    # JSON output
    output_payload = {
        "citation_style": citation_style,
        "max_pages": max_pages,
        "duplicates": duplicates,
        "results": results,
    }
    output_path = output_dir / "_literature_extraction.json"
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(output_payload, f, ensure_ascii=False, indent=2)
    logging.info("JSON 结果已保存：%s", output_path)

    # Markdown report output
    md_path = output_dir / "_literature_report.md"
    md_content = generate_markdown_report(results, lit_dir, skipped, duplicates, citation_style)
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(md_content)
    logging.info("Markdown 报告已保存：%s", md_path)

    # BibTeX output
    if bibtex:
        bib_path = output_dir / "_literature_references.bib"
        with open(bib_path, "w", encoding="utf-8") as f:
            f.write(generate_bibtex(results))
        logging.info("BibTeX 文件已保存：%s", bib_path)


if __name__ == "__main__":
    main()
