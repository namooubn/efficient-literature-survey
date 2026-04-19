#!/usr/bin/env python3
"""Unit and integration tests for efficient-literature-survey.

Run with:
    python -m unittest test_extract_literature_metadata.py
"""

import logging
import os
import sys
import tempfile
import unittest
from pathlib import Path
from unittest.mock import MagicMock, patch

# Ensure sub-package imports resolve when run standalone
_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
if _SCRIPT_DIR not in sys.path:
    sys.path.insert(0, _SCRIPT_DIR)

from cache.manager import file_sha256, load_cache, save_cache
from citation.bibtex import generate_bibtex, _guess_bibtex_type
from citation.engine import generate_citation, _guess_doc_type
from core.helpers import (
    detect_scanned_pdf,
    estimate_pages_from_chars,
    extract_bib_info_from_text,
    find_duplicates,
    get_pdf_read_pages,
    safe_truncate,
    smart_word_count,
    split_chinese_name,
)
from core.logging_config import setup_logging
from extractors.base import make_result_template
from extractors.txt import extract_txt
from extractors.pdf import extract_pdf
from extractors.docx import extract_docx
from extractors.epub import extract_epub

# Optional dependencies for integration tests only
HAS_PYPDF2 = False
HAS_PYTHON_DOCX = False
HAS_EBOOKLIB = False

try:
    import PyPDF2  # noqa: F401
    HAS_PYPDF2 = True
except ImportError:
    pass

try:
    import docx  # noqa: F401
    HAS_PYTHON_DOCX = True
except ImportError:
    pass

try:
    import ebooklib  # noqa: F401
    HAS_EBOOKLIB = True
except ImportError:
    pass


# ---------------------------------------------------------------------------
# Mock-based tests for extractors (no optional deps required)
# ---------------------------------------------------------------------------


class TestExtractPdfMock(unittest.TestCase):
    """Mock tests for extract_pdf that do not require PyPDF2."""

    @patch("extractors.pdf.PdfReader")
    def test_extracts_metadata(self, mock_reader_cls):
        mock_page = MagicMock()
        mock_page.extract_text.return_value = "Page one text content here."
        mock_reader = MagicMock()
        mock_reader.pages = [mock_page]
        mock_reader.metadata = {
            "/Author": "Test Author",
            "/Title": "Test Title",
            "/CreationDate": "D:20230101",
        }
        mock_reader_cls.return_value = mock_reader

        info = extract_pdf("dummy.pdf")
        self.assertEqual(info["author_from_meta"], "Test Author")
        self.assertEqual(info["title_from_meta"], "Test Title")
        self.assertEqual(info["year"], "2023")
        self.assertEqual(info["pages"], 1)
        self.assertFalse(info["is_scanned"])

    @patch("extractors.pdf.PdfReader")
    def test_scanned_detection(self, mock_reader_cls):
        mock_page = MagicMock()
        mock_page.extract_text.return_value = ""
        mock_reader = MagicMock()
        mock_reader.pages = [mock_page] * 10
        mock_reader.metadata = {}
        mock_reader_cls.return_value = mock_reader

        info = extract_pdf("dummy.pdf")
        self.assertTrue(info["is_scanned"])

    @patch("extractors.pdf.PdfReader")
    def test_corrupted_graceful(self, mock_reader_cls):
        mock_reader_cls.side_effect = Exception("corrupted")
        info = extract_pdf("dummy.pdf")
        self.assertIn("corrupted", info["note"].lower())

    @patch("extractors.pdf.PdfReader")
    def test_year_from_filename_fallback(self, mock_reader_cls):
        mock_reader = MagicMock()
        mock_reader.pages = [MagicMock(extract_text=MagicMock(return_value="text"))]
        mock_reader.metadata = {}
        mock_reader_cls.return_value = mock_reader

        info = extract_pdf("paper 2019 final.pdf")
        self.assertEqual(info["year"], "2019")


class TestExtractDocxMock(unittest.TestCase):
    """Mock tests for extract_docx that do not require python-docx."""

    @patch("extractors.docx.Document")
    def test_extracts_metadata(self, mock_doc_cls):
        mock_para = MagicMock()
        mock_para.text = "Hello world this is a test paragraph."
        mock_doc = MagicMock()
        mock_doc.paragraphs = [mock_para]
        mock_doc.core_properties.author = "Docx Author"
        mock_doc.core_properties.title = "Docx Title"
        mock_doc.tables = []
        mock_doc_cls.return_value = mock_doc

        info = extract_docx("dummy.docx")
        self.assertEqual(info["author_from_meta"], "Docx Author")
        self.assertEqual(info["title_from_meta"], "Docx Title")
        self.assertEqual(info["format"], "docx")
        self.assertGreater(info["word_count"], 0)

    @patch("extractors.docx.Document")
    def test_error_handling(self, mock_doc_cls):
        mock_doc_cls.side_effect = Exception("bad docx")
        info = extract_docx("dummy.docx")
        self.assertIn("bad docx", info["note"])


class TestExtractEpubMock(unittest.TestCase):
    """Mock tests for extract_epub that do not require ebooklib."""

    @patch("extractors.epub.epub")
    @patch("extractors.epub.BeautifulSoup")
    @patch("extractors.epub.ebooklib")
    def test_extracts_metadata(self, mock_ebooklib, mock_soup_cls, mock_epub_module):
        mock_item = MagicMock()
        mock_item.get_type.return_value = mock_ebooklib.ITEM_DOCUMENT
        mock_item.get_content.return_value = b"<html><body><p>EPUB text.</p></body></html>"

        mock_book = MagicMock()
        mock_book.get_metadata.side_effect = lambda dc, field: {
            ("DC", "title"): [("Epub Title", {})],
            ("DC", "creator"): [("Epub Author", {})],
        }.get((dc, field), [])
        mock_book.get_items.return_value = [mock_item]
        mock_epub_module.read_epub.return_value = mock_book

        mock_soup = MagicMock()
        mock_soup.get_text.return_value = "EPUB text."
        mock_soup_cls.return_value = mock_soup

        info = extract_epub("dummy.epub")
        self.assertEqual(info["title_from_meta"], "Epub Title")
        self.assertEqual(info["author_from_meta"], "Epub Author")
        self.assertEqual(info["format"], "epub")

    @patch("extractors.epub.epub")
    @patch("extractors.epub.BeautifulSoup")
    @patch("extractors.epub.ebooklib")
    def test_error_handling(self, mock_ebooklib, mock_soup_cls, mock_epub_module):
        mock_epub_module.read_epub.side_effect = Exception("bad epub")
        info = extract_epub("dummy.epub")
        self.assertIn("bad epub", info["note"])


class TestExtractInfoDispatcher(unittest.TestCase):
    """Tests for the top-level extract_info dispatcher."""

    def test_unsupported_format(self):
        from extract_literature_metadata import extract_info
        info = extract_info("file.xyz")
        self.assertIn("Unsupported", info["note"])

    def test_caj_format(self):
        from extract_literature_metadata import extract_info
        info = extract_info("file.caj")
        self.assertEqual(info["format"], "caj")
        self.assertIn("CAJ", info["note"])


# ---------------------------------------------------------------------------
# Core helper tests
# ---------------------------------------------------------------------------


class TestSmartWordCount(unittest.TestCase):
    def test_pure_chinese(self):
        text = "这是一个测试句子，包含多个中文字符。"
        self.assertEqual(smart_word_count(text), 16)

    def test_pure_english(self):
        text = "This is a simple test sentence."
        self.assertEqual(smart_word_count(text), 6)

    def test_mixed_chinese_english(self):
        text = "这是test一个mixed案例case。"
        self.assertEqual(smart_word_count(text), 6)

    def test_empty_string(self):
        self.assertEqual(smart_word_count(""), 0)

    def test_punctuation_only(self):
        self.assertEqual(smart_word_count("！？，。;:"), 1)

    def test_numbers_and_english(self):
        text = "In 2024, there are 3 major updates."
        self.assertEqual(smart_word_count(text), 7)


class TestDetectScannedPdf(unittest.TestCase):
    def test_dense_text_not_scanned(self):
        page_texts = ["a" * 600 for _ in range(10)]
        self.assertFalse(detect_scanned_pdf(page_texts, 10))

    def test_empty_pages_scanned(self):
        page_texts = [""] * 10
        self.assertTrue(detect_scanned_pdf(page_texts, 10))

    def test_cover_image_then_text(self):
        page_texts = [""] + ["a" * 500 for _ in range(9)]
        self.assertFalse(detect_scanned_pdf(page_texts, 10))

    def test_short_doc_safe(self):
        page_texts = ["abc"]
        self.assertFalse(detect_scanned_pdf(page_texts, 1))

    def test_low_density_scanned(self):
        page_texts = ["x" * 8 for _ in range(20)]
        self.assertTrue(detect_scanned_pdf(page_texts, 20))

    def test_no_pages(self):
        self.assertFalse(detect_scanned_pdf([], 0))


class TestEstimatePagesFromChars(unittest.TestCase):
    def test_normal(self):
        self.assertEqual(estimate_pages_from_chars(1000), 2)

    def test_exact_multiple(self):
        self.assertEqual(estimate_pages_from_chars(500), 1)

    def test_zero_fallback(self):
        self.assertEqual(estimate_pages_from_chars(0), 1)

    def test_custom_chars_per_page(self):
        self.assertEqual(estimate_pages_from_chars(1000, chars_per_page=300), 4)


class TestSafeTruncate(unittest.TestCase):
    def test_truncate_long(self):
        text = "a" * 3000
        self.assertEqual(len(safe_truncate(text)), 2000)

    def test_no_truncate_short(self):
        text = "short text"
        self.assertEqual(safe_truncate(text), text)

    def test_empty(self):
        self.assertEqual(safe_truncate(""), "")

    def test_custom_limit(self):
        text = "hello world"
        self.assertEqual(safe_truncate(text, limit=5), "hello")


class TestFileSha256(unittest.TestCase):
    def test_normal_file(self):
        with tempfile.NamedTemporaryFile(mode="w", delete=False, encoding="utf-8") as f:
            f.write("hello world")
            path = f.name
        try:
            h1 = file_sha256(path)
            h2 = file_sha256(path)
            self.assertEqual(len(h1), 64)
            self.assertEqual(h1, h2)
        finally:
            os.unlink(path)

    def test_empty_file(self):
        with tempfile.NamedTemporaryFile(mode="w", delete=False, encoding="utf-8") as f:
            path = f.name
        try:
            h = file_sha256(path)
            self.assertEqual(len(h), 64)
        finally:
            os.unlink(path)

    def test_nonexistent_file(self):
        self.assertEqual(file_sha256("/nonexistent/path/file.txt"), "")

    def test_different_files_different_hash(self):
        with tempfile.NamedTemporaryFile(mode="w", delete=False, encoding="utf-8") as f1:
            f1.write("content A")
            p1 = f1.name
        with tempfile.NamedTemporaryFile(mode="w", delete=False, encoding="utf-8") as f2:
            f2.write("content B")
            p2 = f2.name
        try:
            self.assertNotEqual(file_sha256(p1), file_sha256(p2))
        finally:
            os.unlink(p1)
            os.unlink(p2)


class TestGetPdfReadPages(unittest.TestCase):
    def test_short_doc_reads_all(self):
        self.assertEqual(get_pdf_read_pages(30), list(range(30)))

    def test_exactly_50_reads_all(self):
        self.assertEqual(get_pdf_read_pages(50), list(range(50)))

    def test_long_doc_samples(self):
        pages = get_pdf_read_pages(100)
        self.assertEqual(pages[:15], list(range(15)))
        self.assertEqual(pages[-5:], list(range(95, 100)))
        self.assertEqual(len(pages), 20)

    def test_max_pages_override(self):
        self.assertEqual(get_pdf_read_pages(100, max_pages=5), list(range(5)))

    def test_zero_max_pages_ignored(self):
        self.assertEqual(get_pdf_read_pages(10, max_pages=0), list(range(10)))

    def test_negative_max_pages_ignored(self):
        self.assertEqual(get_pdf_read_pages(10, max_pages=-3), list(range(10)))


class TestFindDuplicates(unittest.TestCase):
    def test_exact_match(self):
        results = [
            {"filename": "a.pdf", "title_from_meta": "Deep Learning"},
            {"filename": "b.pdf", "title_from_meta": "Deep Learning"},
        ]
        dups = find_duplicates(results, threshold=0.90)
        self.assertIn("a.pdf", dups)
        self.assertIn("b.pdf", dups)

    def test_no_match(self):
        results = [
            {"filename": "a.pdf", "title_from_meta": "Deep Learning"},
            {"filename": "b.pdf", "title_from_meta": "Quantum Computing"},
        ]
        dups = find_duplicates(results, threshold=0.90)
        self.assertEqual(dups, {})

    def test_fallback_to_filename(self):
        results = [
            {"filename": "same_title.pdf", "title_from_meta": ""},
            {"filename": "same_title.pdf", "title_from_meta": ""},
        ]
        dups = find_duplicates(results, threshold=0.90)
        self.assertIn("same_title.pdf", dups)


class TestGenerateCitation(unittest.TestCase):
    def test_gb7714_single_author(self):
        meta = {"author_from_meta": "张三", "title_from_meta": "人工智能导论", "filename": "x.pdf", "format": "pdf"}
        cite = generate_citation(meta, "gb7714")
        self.assertIn("张三", cite)
        self.assertIn("人工智能导论", cite)

    def test_gb7714_multiple_authors(self):
        meta = {"author_from_meta": "张三, 李四, 王五, 赵六", "title_from_meta": "深度学习", "filename": "x.pdf", "format": "pdf"}
        cite = generate_citation(meta, "gb7714")
        self.assertIn("张三", cite)
        self.assertIn("等", cite)

    def test_gb7714_with_journal(self):
        meta = {
            "author_from_meta": "张三",
            "title_from_meta": "测试",
            "filename": "x.pdf",
            "format": "pdf",
            "year": "2023",
            "journal": "计算机学报",
            "volume": "45",
            "issue": "3",
            "page_range": "120-130",
        }
        cite = generate_citation(meta, "gb7714")
        self.assertIn("计算机学报", cite)
        self.assertIn("2023", cite)
        self.assertIn("45(3)", cite)
        self.assertIn("120-130", cite)
        self.assertNotIn("信息不完整", cite)

    def test_gb7714_incomplete(self):
        meta = {"author_from_meta": "张三", "title_from_meta": "测试", "filename": "x.pdf", "format": "pdf"}
        cite = generate_citation(meta, "gb7714")
        self.assertIn("信息不完整", cite)

    def test_apa(self):
        meta = {"author_from_meta": "John Smith", "title_from_meta": "AI Survey", "filename": "x.pdf", "format": "pdf"}
        cite = generate_citation(meta, "apa")
        self.assertIn("Smith", cite)
        self.assertIn("AI Survey", cite)

    def test_apa_with_journal(self):
        meta = {
            "author_from_meta": "John Smith",
            "title_from_meta": "AI Survey",
            "filename": "x.pdf",
            "format": "pdf",
            "year": "2023",
            "journal": "Nature",
            "volume": "10",
            "issue": "2",
            "page_range": "1-10",
        }
        cite = generate_citation(meta, "apa")
        self.assertIn("Smith", cite)
        self.assertIn("(2023)", cite)
        self.assertIn("*Nature*", cite)
        self.assertIn("*10*(2)", cite)

    def test_mla(self):
        meta = {"author_from_meta": "John Smith", "title_from_meta": "AI Survey", "filename": "x.pdf", "format": "pdf"}
        cite = generate_citation(meta, "mla")
        self.assertIn('"AI Survey."', cite)
        self.assertIn("Smith", cite)

    def test_numbered(self):
        meta = {"author_from_meta": "", "title_from_meta": "Title", "filename": "x.pdf", "format": "pdf", "_citation_number": 5, "year": "2023", "publisher": "Test Press"}
        self.assertEqual(generate_citation(meta, "numbered"), "[5] Title, Test Press, 2023.")

    def test_no_author(self):
        meta = {"author_from_meta": "", "title_from_meta": "", "filename": "unknown.pdf", "format": "pdf"}
        cite = generate_citation(meta, "gb7714")
        self.assertIn("unknown.pdf", cite)

    def test_apa_compound_surname(self):
        meta = {"author_from_meta": "欧阳明", "title_from_meta": "测试", "filename": "x.pdf", "format": "pdf"}
        cite = generate_citation(meta, "apa")
        self.assertIn("欧阳", cite)

    def test_mla_compound_surname(self):
        meta = {"author_from_meta": "欧阳明", "title_from_meta": "测试", "filename": "x.pdf", "format": "pdf"}
        cite = generate_citation(meta, "mla")
        self.assertIn("欧阳, 明", cite)


class TestSplitChineseName(unittest.TestCase):
    def test_single_surname(self):
        self.assertEqual(split_chinese_name("张三"), ("张", "三"))

    def test_compound_surname(self):
        self.assertEqual(split_chinese_name("欧阳明"), ("欧阳", "明"))

    def test_compound_surname_long(self):
        self.assertEqual(split_chinese_name("司马相如"), ("司马", "相如"))

    def test_empty(self):
        self.assertEqual(split_chinese_name(""), ("", ""))

    def test_with_spaces(self):
        self.assertEqual(split_chinese_name("  张三  "), ("张", "三"))


class TestExtractBibInfoFromText(unittest.TestCase):
    def test_doi_extraction(self):
        text = "This paper is published with DOI: 10.1234/example.5678"
        info = extract_bib_info_from_text(text)
        self.assertEqual(info["doi"], "10.1234/example.5678")

    def test_volume_issue_english(self):
        text = "Journal of AI, Vol. 12, No. 3, pp. 45-67"
        info = extract_bib_info_from_text(text)
        self.assertEqual(info["volume"], "12")
        self.assertEqual(info["issue"], "3")
        self.assertEqual(info["page_range"], "45-67")

    def test_chinese_volume_issue(self):
        text = "发表于《计算机学报》第 15 卷第 2 期，第 100-120 页"
        info = extract_bib_info_from_text(text)
        self.assertEqual(info["volume"], "15")
        self.assertEqual(info["issue"], "2")
        self.assertEqual(info["page_range"], "100-120")
        self.assertEqual(info["journal"], "计算机学报")

    def test_no_match(self):
        info = extract_bib_info_from_text("")
        self.assertEqual(info["doi"], "")
        self.assertEqual(info["journal"], "")

    def test_journal_book_title(self):
        text = "Published in 《自然语言处理》杂志"
        info = extract_bib_info_from_text(text)
        self.assertEqual(info["journal"], "自然语言处理")


class TestGuessDocType(unittest.TestCase):
    def test_short_pdf_is_journal(self):
        self.assertEqual(_guess_doc_type({"format": "pdf", "pages": 10}), "J")

    def test_long_pdf_is_monograph(self):
        self.assertEqual(_guess_doc_type({"format": "pdf", "pages": 250}), "M")

    def test_epub_is_monograph(self):
        self.assertEqual(_guess_doc_type({"format": "epub", "pages": 50}), "M")


class TestGuessBibtexType(unittest.TestCase):
    def test_with_journal(self):
        self.assertEqual(_guess_bibtex_type({"journal": "Nature", "pages": 10}), "article")

    def test_long_no_journal(self):
        self.assertEqual(_guess_bibtex_type({"journal": "", "pages": 300}), "book")

    def test_default_misc(self):
        self.assertEqual(_guess_bibtex_type({"journal": "", "pages": 10}), "misc")


class TestGenerateBibtex(unittest.TestCase):
    def test_single_entry(self):
        results = [
            {
                "author_from_meta": "张三",
                "title_from_meta": "测试标题",
                "filename": "test.pdf",
                "format": "pdf",
                "year": "2023",
                "journal": "计算机学报",
                "volume": "45",
                "issue": "3",
                "page_range": "100-110",
                "doi": "10.1234/test",
            }
        ]
        bib = generate_bibtex(results)
        self.assertIn("@article", bib)
        self.assertIn("测试标题", bib)
        self.assertIn("张三", bib)
        self.assertIn("2023", bib)
        self.assertIn("计算机学报", bib)
        self.assertIn("10.1234/test", bib)

    def test_incomplete_entry_has_note(self):
        results = [
            {
                "author_from_meta": "",
                "title_from_meta": "",
                "filename": "unknown.pdf",
                "format": "pdf",
            }
        ]
        bib = generate_bibtex(results)
        self.assertIn("文献信息不完整", bib)

    def test_multiple_entries(self):
        results = [
            {"author_from_meta": "A", "title_from_meta": "T1", "filename": "a.pdf", "format": "pdf", "year": "2021"},
            {"author_from_meta": "B", "title_from_meta": "T2", "filename": "b.pdf", "format": "pdf", "year": "2022"},
        ]
        bib = generate_bibtex(results)
        self.assertEqual(bib.count("@"), 2)


class TestSetupLogging(unittest.TestCase):
    def test_verbose_sets_debug(self):
        setup_logging(verbose=True, quiet=False)
        self.assertEqual(logging.getLogger().level, logging.DEBUG)

    def test_quiet_sets_warning(self):
        setup_logging(verbose=False, quiet=True)
        self.assertEqual(logging.getLogger().level, logging.WARNING)

    def test_default_sets_info(self):
        setup_logging(verbose=False, quiet=False)
        self.assertEqual(logging.getLogger().level, logging.INFO)


class TestCacheManager(unittest.TestCase):
    """Tests for cache manager keyed by relative path."""

    def test_roundtrip(self):
        with tempfile.TemporaryDirectory() as td:
            path = Path(td) / "cache.json"
            data = {"files": {"subdir/paper.pdf": {"sha256": "abc123", "result": {"pages": 10}}}}
            save_cache(path, data["files"])
            loaded = load_cache(path)
            self.assertEqual(loaded["subdir/paper.pdf"]["sha256"], "abc123")

    def test_missing_file_returns_empty(self):
        loaded = load_cache(Path("/nonexistent/cache.json"))
        self.assertEqual(loaded, {})

    def test_version_in_payload(self):
        with tempfile.TemporaryDirectory() as td:
            path = Path(td) / "cache.json"
            save_cache(path, {})
            with open(path, "r", encoding="utf-8") as f:
                import json
                payload = json.load(f)
            self.assertEqual(payload["version"], 2)
            self.assertIn("generated_at", payload)


class TestExtractTxt(unittest.TestCase):
    """Integration tests for extract_txt (no extra deps)."""

    def test_utf8_file(self):
        with tempfile.NamedTemporaryFile(mode="w", delete=False, encoding="utf-8", suffix=".txt") as f:
            f.write("Hello world\nSecond line")
            path = f.name
        try:
            info = extract_txt(path)
            self.assertEqual(info["word_count"], 4)
            self.assertEqual(info["text_chars"], 23)
            self.assertEqual(info["format"], "txt")
        finally:
            os.unlink(path)

    def test_gbk_file(self):
        with tempfile.NamedTemporaryFile(mode="w", delete=False, encoding="gbk", suffix=".txt") as f:
            f.write("中文内容测试")
            path = f.name
        try:
            info = extract_txt(path)
            self.assertEqual(info["word_count"], 6)
            self.assertEqual(info["format"], "txt")
        finally:
            os.unlink(path)

    def test_empty_file(self):
        with tempfile.NamedTemporaryFile(mode="w", delete=False, encoding="utf-8", suffix=".txt") as f:
            path = f.name
        try:
            info = extract_txt(path)
            self.assertEqual(info["word_count"], 0)
            self.assertEqual(info["text_chars"], 0)
        finally:
            os.unlink(path)

    def test_md_file(self):
        with tempfile.NamedTemporaryFile(mode="w", delete=False, encoding="utf-8", suffix=".md") as f:
            f.write("# Title\n\nBody text here.")
            path = f.name
        try:
            info = extract_txt(path)
            self.assertEqual(info["format"], "md")
            self.assertIn("Title", info["first_page_text"])
        finally:
            os.unlink(path)


# ---------------------------------------------------------------------------
# Integration tests requiring optional dependencies
# ---------------------------------------------------------------------------


@unittest.skipUnless(HAS_PYPDF2, "PyPDF2 not installed")
class TestExtractPdfIntegration(unittest.TestCase):
    """Integration tests for extract_pdf."""

    def _make_blank_pdf(self, page_count: int) -> str:
        from PyPDF2 import PdfWriter
        writer = PdfWriter()
        for _ in range(page_count):
            writer.add_blank_page(width=612, height=792)
        fd, path = tempfile.mkstemp(suffix=".pdf")
        try:
            with os.fdopen(fd, "wb") as f:
                writer.write(f)
        except Exception:
            os.close(fd)
            raise
        return path

    def test_normal_pdf_pages(self):
        path = self._make_blank_pdf(10)
        try:
            info = extract_pdf(path)
            self.assertEqual(info["pages"], 10)
            self.assertEqual(info["format"], "pdf")
        finally:
            os.unlink(path)

    def test_short_pdf_reads_all_pages(self):
        path = self._make_blank_pdf(20)
        try:
            info = extract_pdf(path)
            self.assertEqual(info["pages"], 20)
            self.assertTrue(info["is_scanned"])
        finally:
            os.unlink(path)

    def test_long_pdf_samples_pages(self):
        path = self._make_blank_pdf(100)
        try:
            info = extract_pdf(path)
            self.assertEqual(info["pages"], 100)
            self.assertEqual(info["text_chars"], 0)
        finally:
            os.unlink(path)

    def test_corrupted_pdf_graceful(self):
        fd, path = tempfile.mkstemp(suffix=".pdf")
        try:
            with os.fdopen(fd, "w", encoding="utf-8") as f:
                f.write("this is not a pdf")
            info = extract_pdf(path)
            self.assertIn("error", info["note"].lower())
        finally:
            os.unlink(path)

    def test_max_pages_override(self):
        path = self._make_blank_pdf(30)
        try:
            info = extract_pdf(path, max_pages=5)
            self.assertEqual(info["pages"], 30)
            self.assertEqual(info["text_chars"], 0)
        finally:
            os.unlink(path)


@unittest.skipUnless(HAS_PYTHON_DOCX, "python-docx not installed")
class TestExtractDocxIntegration(unittest.TestCase):
    """Integration tests for extract_docx."""

    def test_normal_docx(self):
        from docx import Document
        doc = Document()
        doc.add_paragraph("Hello world")
        doc.add_paragraph("Second paragraph with more text")

        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        try:
            doc.save(path)
            info = extract_docx(path)
            self.assertEqual(info["format"], "docx")
            self.assertGreater(info["word_count"], 0)
            self.assertIn("Hello", info["first_page_text"])
        finally:
            os.unlink(path)

    def test_empty_docx(self):
        from docx import Document
        doc = Document()
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        try:
            doc.save(path)
            info = extract_docx(path)
            self.assertEqual(info["format"], "docx")
            self.assertEqual(info["word_count"], 0)
        finally:
            os.unlink(path)

    def test_docx_with_author_title(self):
        from docx import Document
        doc = Document()
        doc.add_paragraph("Content here")
        doc.core_properties.author = "Test Author"
        doc.core_properties.title = "Test Title"

        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        try:
            doc.save(path)
            info = extract_docx(path)
            self.assertEqual(info["author_from_meta"], "Test Author")
            self.assertEqual(info["title_from_meta"], "Test Title")
        finally:
            os.unlink(path)


@unittest.skipUnless(HAS_EBOOKLIB, "ebooklib not installed")
class TestExtractEpubIntegration(unittest.TestCase):
    """Integration tests for extract_epub."""

    def test_normal_epub(self):
        from ebooklib import epub
        book = epub.EpubBook()
        book.set_identifier("test-id")
        book.set_title("Test Book")
        book.set_language("en")
        book.add_author("Test Author")

        c1 = epub.EpubHtml(title="Chapter 1", file_name="chap_01.xhtml", lang="en")
        c1.content = "<html><body><p>Hello world from EPUB.</p></body></html>"
        book.add_item(c1)
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        book.spine = [c1]

        fd, path = tempfile.mkstemp(suffix=".epub")
        os.close(fd)
        try:
            epub.write_epub(path, book)
            info = extract_epub(path)
            self.assertEqual(info["format"], "epub")
            self.assertEqual(info["title_from_meta"], "Test Book")
            self.assertEqual(info["author_from_meta"], "Test Author")
            self.assertIn("Hello", info["first_page_text"])
        finally:
            os.unlink(path)


if __name__ == "__main__":
    unittest.main()
