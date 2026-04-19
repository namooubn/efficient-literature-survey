"""DOCX metadata and text extraction."""

from .base import make_result_template
from core.config import DOCX_SAMPLE_CHARS, DOCX_TABLE_SAMPLE_COUNT
from core.helpers import (
    estimate_pages_from_chars,
    extract_bib_info_from_text,
    safe_truncate,
    smart_word_count,
)

try:
    from docx import Document
except ImportError:  # pragma: no cover
    Document = None  # type: ignore[misc, assignment]


def extract_docx(docx_path: str) -> dict:
    """Extract metadata and text sample from a DOCX file."""
    result = make_result_template(docx_path)
    result["format"] = "docx"

    if Document is None:
        result["note"] = "python-docx not installed. Run: pip install python-docx"
        return result

    try:
        doc = Document(docx_path)

        core_props = doc.core_properties
        result["author_from_meta"] = core_props.author or ""
        result["title_from_meta"] = core_props.title or ""

        paragraphs = []
        char_count = 0
        for p in doc.paragraphs:
            paragraphs.append(p)
            char_count += len(p.text)
            if char_count >= DOCX_SAMPLE_CHARS:
                break
        full_text = "\n".join(p.text for p in paragraphs)
        result["text_chars"] = len(full_text)
        result["word_count"] = smart_word_count(full_text)
        result["first_page_text"] = safe_truncate(full_text)
        result["pages"] = estimate_pages_from_chars(result["text_chars"])

        table_texts = []
        for table in doc.tables[:DOCX_TABLE_SAMPLE_COUNT]:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        table_texts.append(cell.text.strip())
        if table_texts:
            table_full = "\n".join(table_texts)
            result["word_count"] += smart_word_count(table_full)

        bib_info = extract_bib_info_from_text(full_text)
        result.update(bib_info)

    except Exception as e:
        result["note"] = f"DOCX extraction error: {e}"

    return result
